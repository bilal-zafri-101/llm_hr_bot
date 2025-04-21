import fitz  # PyMuPDF
import docx
import io
from typing import Union

def extract_text_from_pdf(file: Union[io.BytesIO, str]) -> str:
    try:
        if isinstance(file, io.BytesIO):
            doc = fitz.open(stream=file.read(), filetype="pdf")
        else:
            doc = fitz.open(file)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        return f"[ERROR] Failed to parse PDF: {e}"

def extract_text_from_docx(file: Union[io.BytesIO, str]) -> str:
    try:
        if isinstance(file, io.BytesIO):
            doc = docx.Document(file)
        else:
            with open(file, 'rb') as f:
                doc = docx.Document(f)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        return f"[ERROR] Failed to parse DOCX: {e}"

def parse_resume(file: Union[io.BytesIO, str], filename: str = "") -> str:
    filename = filename.lower()
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file)
    else:
        return "[ERROR] Unsupported file format. Use PDF or DOCX."
